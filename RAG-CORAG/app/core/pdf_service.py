"""
pdf_service.py
--------------
Xử lý tài liệu upload: PDF/DOCX, tách chunks, tạo FAISS vector store.
"""

import os
import uuid
import logging
from contextlib import contextmanager
from pathlib import Path

from pypdf import PdfReader, errors as pypdf_errors
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .dependencies import get_embeddings
from .exceptions import DocumentProcessingError, DocumentValidationError

logger = logging.getLogger(__name__)

UPLOAD_DIR = "/tmp/document_uploads"
MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

PDF_MAGIC_BYTES = b"%PDF"
DOCX_MAGIC_BYTES = b"PK\x03\x04"
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

os.makedirs(UPLOAD_DIR, exist_ok=True)


def _get_file_extension(file_obj) -> str:
    return Path(getattr(file_obj, "name", "")).suffix.lower()


def validate_upload(file_obj) -> str:
    """
    Kiểm tra file trước khi lưu tạm:
      1. Phần mở rộng phải là .pdf hoặc .docx
      2. Kích thước phải < 50 MB
      3. Magic bytes phải khớp định dạng
    Returns extension hợp lệ.
    """
    name = getattr(file_obj, "name", "")
    extension = _get_file_extension(file_obj)
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentValidationError(
            f"File '{name}' không hợp lệ. "
            "Chỉ chấp nhận file có đuôi .pdf hoặc .docx"
        )

    size = getattr(file_obj, "size", None)
    if size is None:
        file_obj.seek(0, 2)
        size = file_obj.tell()
        file_obj.seek(0)

    if size > MAX_FILE_SIZE_BYTES:
        size_mb = size / (1024 * 1024)
        raise DocumentValidationError(
            f"File quá lớn: {size_mb:.1f} MB (giới hạn {MAX_FILE_SIZE_MB} MB). "
            "File lớn sẽ làm chậm quá trình xử lý đáng kể. "
            "Hãy tách tài liệu thành các phần nhỏ hơn."
        )

    header = file_obj.read(4)
    file_obj.seek(0)

    if extension == ".pdf" and header != PDF_MAGIC_BYTES:
        raise DocumentValidationError(
            "File không phải PDF hợp lệ (magic bytes không khớp). "
            "File có thể bị đổi tên hoặc bị corrupt."
        )

    if extension == ".docx" and header != DOCX_MAGIC_BYTES:
        raise DocumentValidationError(
            "File không phải DOCX hợp lệ (magic bytes không khớp). "
            "File có thể bị đổi tên hoặc bị corrupt."
        )

    return extension


@contextmanager
def temp_document(file_obj):
    """
    Validate → lưu vào đường dẫn tạm duy nhất (UUID) → tự cleanup.
    """
    extension = validate_upload(file_obj)

    path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}{extension}")
    try:
        with open(path, "wb+") as f:
            for chunk in file_obj.chunks():
                f.write(chunk)
        yield path
    finally:
        if os.path.exists(path):
            os.remove(path)
            logger.debug(f"Cleaned up temp file: {path}")


def load_pdf(file_path: str) -> str:
    """
    Trích xuất toàn bộ text từ PDF.
    Raises:
        PDFValidationError  — PDF bị mã hóa (cần password)
        PDFProcessingError  — PDF là ảnh scan hoặc bị corrupt
    """
    try:
        reader = PdfReader(file_path)
    except pypdf_errors.PdfReadError as e:
        raise DocumentProcessingError(
            f"Không đọc được PDF (file có thể bị corrupt): {e}"
        ) from e

    if reader.is_encrypted:
        raise DocumentValidationError(
            "PDF được bảo vệ bằng mật khẩu. "
            "Hãy gỡ mật khẩu trước khi upload."
        )

    total_pages = len(reader.pages)
    pages_text, empty_pages = [], []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(text)
        else:
            empty_pages.append(i + 1)

    if empty_pages:
        logger.warning(
            f"{len(empty_pages)}/{total_pages} trang không có text "
            f"(trang: {empty_pages[:10]}{'...' if len(empty_pages) > 10 else ''}). "
            "Các trang này có thể là ảnh scan."
        )

    full_text = "\n".join(pages_text).strip()
    if not full_text:
        raise DocumentProcessingError(
            f"PDF có {total_pages} trang nhưng không trích xuất được text nào. "
            "File này có thể là ảnh scan — hãy chạy OCR trước (vd: Adobe Acrobat, "
            "tesseract, hoặc các công cụ online)."
        )

    logger.info(
        f"Đọc PDF thành công: {total_pages} trang, "
        f"{len(pages_text)} trang có text, "
        f"{len(full_text):,} ký tự."
    )
    return full_text


def load_docx(file_path: str) -> str:
    """
    Trích xuất toàn bộ text từ DOCX.
    """
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as e:
        raise DocumentProcessingError(
            "Thiếu thư viện python-docx để đọc file DOCX. "
            "Hãy cài bằng lệnh: pip install python-docx"
        ) from e

    try:
        document = DocxDocument(file_path)
    except Exception as e:
        raise DocumentProcessingError(
            f"Không đọc được DOCX (file có thể bị corrupt): {e}"
        ) from e

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + tables).strip()
    if not full_text:
        raise DocumentProcessingError(
            "DOCX không trích xuất được text nào. "
            "Hãy kiểm tra tài liệu có nội dung văn bản hay không."
        )

    logger.info(
        f"Đọc DOCX thành công: {len(paragraphs)} đoạn văn, "
        f"{len(tables)} dòng bảng, {len(full_text):,} ký tự."
    )
    return full_text


def load_document(file_path: str) -> str:
    extension = Path(file_path).suffix.lower()
    if extension == ".pdf":
        return load_pdf(file_path)
    if extension == ".docx":
        return load_docx(file_path)
    raise DocumentValidationError("Định dạng file chưa được hỗ trợ.")


def create_vector_store(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
) -> FAISS:
    """
    Tách text thành chunks → embed → lưu vào FAISS.
    Raises DocumentProcessingError nếu không tách được chunk.
    """
    if chunk_size <= 0:
        raise DocumentValidationError("chunk_size phải lớn hơn 0.")
    if chunk_overlap < 0:
        raise DocumentValidationError("chunk_overlap không được âm.")
    if chunk_overlap >= chunk_size:
        raise DocumentValidationError("chunk_overlap phải nhỏ hơn chunk_size.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " "],
    )
    chunks = splitter.split_text(text)
    if not chunks:
        raise DocumentProcessingError("Không tách được chunk từ nội dung tài liệu.")

    logger.info(
        f"Tạo vector store từ {len(chunks)} chunks "
        f"(chunk_size={chunk_size}, chunk_overlap={chunk_overlap})..."
    )
    embeddings = get_embeddings()
    return FAISS.from_texts(chunks, embeddings)


# ------------------------------------------------------------------ #
#  Lấy chunks để summarize (đầu + cuối tài liệu)                     #
# ------------------------------------------------------------------ #
def get_summary_chunks(vector_store: FAISS, top_k: int = 5) -> str:
    """
    Dùng similarity search với nhiều query đa dạng thay vì
    một query cố định, để tóm tắt bao quát hơn.
    """
    summary_queries = [
        "nội dung chính của tài liệu",
        "kết luận và kết quả",
        "giới thiệu và mục tiêu",
    ]
    seen, chunks = set(), []
    for q in summary_queries:
        for doc in vector_store.similarity_search(q, k=top_k):
            if doc.page_content not in seen:
                seen.add(doc.page_content)
                chunks.append(doc.page_content)

    return "\n\n".join(chunks)
